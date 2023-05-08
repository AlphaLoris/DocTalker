import os
import tkinter as tk
from tkinter import filedialog
from docx import Document

"""
This strips all formatting, hyperlinks, tables, images, drawings, comments, headers, footers, and styles from a .docx
file. This approach (stripping out stuff that isn't valuable vs pulling out the stuff that is valuable)
is probably a less useful way to approach this problem, so I'm not sure if I'll use this script in the future.
"""


def strip_elements(docx_path, output_path):
    doc = Document(docx_path)

    for paragraph in doc.paragraphs:
        # Remove text formatting
        for run in paragraph.runs:
            run.font.name = None
            run.font.size = None
            run.font.bold = None
            run.font.italic = None
            run.font.underline = None

        # Remove hyperlinks
        for hlink in paragraph._element.xpath(".//w:hyperlink"):
            hlink.getparent().remove(hlink)

    # Remove tables
    for table in doc.tables:
        table._element.getparent().remove(table._element)

    # Remove images and drawings
    for rel in list(doc.part.rels.values()):
        if rel.reltype.endswith("image") or rel.reltype.endswith("drawing"):
            doc.part.drop_rel(rel.rId)

    # Remove comments
    for comment in doc._element.xpath(".//w:commentRangeStart | .//w:commentRangeEnd | .//w:commentReference"):
        comment.getparent().remove(comment)

    # Remove headers and footers
    for section in doc.sections:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

    # Remove styles
    for style in doc.styles:
        if style.builtin:
            style.hidden = True

    doc.save(output_path)


def open_file_browser():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    return file_path


def main():
    docx_path = open_file_browser()
    if not docx_path:
        print("No file selected. Exiting.")
        return

    output_path = os.path.splitext(docx_path)[0] + "-Simple.docx"
    strip_elements(docx_path, output_path)
    print(f"Modified file saved as {output_path}")


if __name__ == "__main__":
    main()
