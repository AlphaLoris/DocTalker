import docx
from bs4 import BeautifulSoup
import spacy

nlp = spacy.load("en_core_web_sm")

# TODO: Look at the metadata and consider the semantic tagging as elements to be leveraged.

"""
This script processes a Microsoft Word document (.docx) and extracts text, formatting, and semantic information.

** It creates very small chunks of text (sentences) and performs semantic tagging on them using the spacy library. **

It does the following steps:

    Imports required libraries (docx, BeautifulSoup, spacy).

    Loads the 'en_core_web_sm' model for the spacy library.

    Defines the process_word_document function, which takes a file path as its input, and:

    a. Reads the document using the python-docx library.

    b. Extracts text, paragraph style, and inline formatting information (bold, italic, underline)
    from each paragraph in the document, and stores it in a list called extracted_data.

    c. Converts tables and lists in the document into a simple text format and appends them to the extracted_data list.
    Note that footnotes, endnotes, hyperlinks, and multimedia extraction are not supported by the python-docx library.

    d. Chunks the text into sentences using the spacy library.

    e. Performs semantic tagging on each chunk (sentence) by extracting named entities and their labels
    using the spacy library. The results are stored in a list called tagged_chunks.

    f. Indexes the tagged_chunks using an external search engine (e.g., Elasticsearch, Solr).
    This step is not implemented in the code and should be done separately, depending on the search engine used.

    g. Returns the tagged_chunks list.

    When executed as a standalone script, it reads a specified Word document, processes it using the 
    process_word_document function, and prints the processed data to the console.
    """


def process_word_document(file_path):
    doc = docx.Document(file_path)

    # Text extraction and preservation of hierarchical structures and inline formatting
    extracted_data = []
    for paragraph in doc.paragraphs:
        paragraph_data = {
            'text': paragraph.text,
            'style': paragraph.style.name,
            'inline_formatting': []
        }
        for run in paragraph.runs:
            if run.bold or run.italic or run.underline:
                inline_format = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline
                }
                paragraph_data['inline_formatting'].append(inline_format)
        extracted_data.append(paragraph_data)

    # Tables and lists conversion (assuming simple tables)
    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text for cell in row.cells]
            extracted_data.append({'table_row': cells_text})

    # Footnotes, endnotes, hyperlinks, and multimedia extraction is not supported by python-docx
    # For this, you may need to use more advanced libraries or custom code

    # Chunking
    chunks = []
    for data in extracted_data:
        doc = nlp(data['text'])
        for sent in doc.sents:
            chunks.append(sent.text)

    # Semantic tagging
    tagged_chunks = []
    for chunk in chunks:
        doc = nlp(chunk)
        entities = [(ent.text, ent.label_) for ent in doc.ents]
        tagged_chunks.append({
            'chunk': chunk,
            'entities': entities
        })

    # Indexing
    # Assuming you have already set up Elasticsearch or Solr, index the tagged_chunks
    # For example, with Elasticsearch, you would use the elasticsearch-py library
    # and create an index using the client.index() method

    return tagged_chunks


if __name__ == "__main__":
    file_path = r"C:\Users\glenn\OneDrive\Documents\Glenn's Docs\Linklings\WIP\Experiments\Copy of Linklings Handbook for Chairs and Admins.docx"

    processed_data = process_word_document(file_path)
    print(processed_data)
